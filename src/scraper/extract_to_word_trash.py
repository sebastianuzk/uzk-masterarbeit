"""
Extrahiere HTML-Inhalt aus content_database und erstelle Word-Dokument
"""
import sqlite3
import gzip
from pathlib import Path
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Lade Dokument aus content_database
db_path = Path('data/content_database.db')
conn = sqlite3.connect(db_path)
cursor = conn.cursor()

# Zeige verfügbare Tabellen
cursor.execute("SELECT name FROM sqlite_master WHERE type='table'")
tables = cursor.fetchall()
print('Verfügbare Tabellen:')
for table in tables:
    print(f'  - {table[0]}')

# Prüfe Schema
if tables:
    table_name = tables[0][0]
    cursor.execute(f'PRAGMA table_info({table_name})')
    columns = cursor.fetchall()
    print(f'\nSpalten in {table_name}:')
    for col in columns:
        print(f'  - {col[1]} ({col[2]})')

# Suche URL
url = 'https://wiso.uni-koeln.de/de/praxis/veranstaltungen/veranstaltungen-vergangener-semester/wise-2024-25'
cursor.execute(f"SELECT * FROM {table_name} WHERE url = ?", (url,))
result = cursor.fetchone()

if not result:
    print(f'\n❌ URL nicht gefunden: {url}')
    # Suche ähnliche URLs
    cursor.execute(f"SELECT url FROM {table_name} WHERE url LIKE '%wise-2024-25%'")
    similar = cursor.fetchall()
    if similar:
        print('\nÄhnliche URLs:')
        for s in similar:
            print(f'  - {s[0]}')
    conn.close()
    exit(1)

print(f'\n✅ Dokument gefunden!')

# Finde Content-Spalte
content_idx = None
for i, col in enumerate(columns):
    if 'content' in col[1].lower():
        content_idx = i
        break

if content_idx is None:
    print('❌ Keine Content-Spalte gefunden')
    conn.close()
    exit(1)

# Dekomprimiere Content
html = gzip.decompress(result[content_idx]).decode('utf-8')
soup = BeautifulSoup(html, 'html.parser')

# Entferne Script/Style/Nav
for element in soup(['script', 'style', 'head', 'meta', 'link', 'noscript', 'iframe', 'nav', 'footer', 'aside']):
    element.decompose()

# Erstelle Word-Dokument
doc = Document()

# Titel
title = doc.add_heading('WiSe 2024/25 - Veranstaltungen vergangener Semester', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# URL als Untertitel
url_para = doc.add_paragraph(url)
url_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
url_para.runs[0].font.size = Pt(9)
url_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph()  # Leerzeile

# Extrahiere strukturierten Content
for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'blockquote']):
    if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
        level = int(element.name[1])
        text = element.get_text(strip=True)
        if text:
            doc.add_heading(text, level)
    
    elif element.name == 'p':
        text = element.get_text(strip=True)
        if text:
            doc.add_paragraph(text)
    
    elif element.name in ['ul', 'ol']:
        for li in element.find_all('li', recursive=False):
            text = li.get_text(strip=True)
            if text:
                doc.add_paragraph(text, style='List Bullet' if element.name == 'ul' else 'List Number')
    
    elif element.name == 'blockquote':
        text = element.get_text(strip=True)
        if text:
            para = doc.add_paragraph(text)
            para.style = 'Quote'

# Speichere Word-Dokument
output_path = Path('data/deduplication/wise_2024-25_veranstaltungen.docx')
doc.save(output_path)

print(f'\n✅ Word-Dokument erstellt: {output_path}')
print(f'Anzahl Absätze: {len(doc.paragraphs)}')

conn.close()
